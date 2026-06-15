# -*- coding: utf-8 -*-
"""Convert paper_NAO_delivery.md to a Word .docx file with IEEE-like formatting."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = r"c:\Users\DELL\Documents\robot_Nao\paper_NAO_delivery.md"
DST = r"c:\Users\DELL\Documents\robot_Nao\paper_NAO_delivery.docx"


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def find_matching_brace(s, start):
    """Given s[start] == '{', return index of matching '}'. -1 if not found."""
    depth = 1
    i = start + 1
    while i < len(s):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return -1


def expand_frac(s):
    """Replace \\frac{a}{b} with (a)/(b), supporting nested braces."""
    while '\\frac{' in s:
        idx = s.find('\\frac{')
        a_start = idx + 6
        a_end = find_matching_brace(s, a_start - 1)
        if a_end == -1:
            break
        if a_end + 1 >= len(s) or s[a_end + 1] != '{':
            break
        b_start = a_end + 2
        b_end = find_matching_brace(s, b_start - 1)
        if b_end == -1:
            break
        a = s[a_start:a_end]
        b = s[b_start:b_end]
        s = s[:idx] + f'({a}) / ({b})' + s[b_end + 1:]
    return s


def latex_to_unicode(s):
    """Convert simplified LaTeX commands to Unicode (no sub/sup handling)."""
    s = expand_frac(s)
    # Cases environment: \begin{cases} a & b \\ c & d \end{cases}
    s = re.sub(r'\\begin\{cases\}', '{ ', s)
    s = re.sub(r'\\end\{cases\}', '', s)
    # \text{...} -> ...
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)
    # \mathrm{...} -> ...
    s = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', s)
    pairs = [
        (r'\sqrt{2}', '√2'),
        (r'\Delta', 'Δ'),
        (r'\delta', 'δ'),
        (r'\epsilon', 'ε'),
        (r'\varepsilon', 'ε'),
        (r'\pi', 'π'),
        (r'\Sigma', 'Σ'),
        (r'\sigma', 'σ'),
        (r'\alpha', 'α'),
        (r'\beta', 'β'),
        (r'\gamma', 'γ'),
        (r'\theta', 'θ'),
        (r'\lambda', 'λ'),
        (r'\mu', 'μ'),
        (r'\ldots', '…'),
        (r'\dots', '…'),
        (r'\cdots', '⋯'),
        (r'\leq', '≤'),
        (r'\le', '≤'),
        (r'\geq', '≥'),
        (r'\ge', '≥'),
        (r'\neq', '≠'),
        (r'\ne', '≠'),
        (r'\approx', '≈'),
        (r'\equiv', '≡'),
        (r'\in', '∈'),
        (r'\notin', '∉'),
        (r'\subset', '⊂'),
        (r'\supset', '⊃'),
        (r'\cup', '∪'),
        (r'\cap', '∩'),
        (r'\emptyset', '∅'),
        (r'\infty', '∞'),
        (r'\times', '×'),
        (r'\cdot', '·'),
        (r'\pm', '±'),
        (r'\to', '→'),
        (r'\rightarrow', '→'),
        (r'\leftarrow', '←'),
        (r'\Rightarrow', '⇒'),
        (r'\min', 'min'),
        (r'\max', 'max'),
        (r'\arg', 'arg'),
        (r'\sum', 'Σ'),
        (r'\prod', '∏'),
        (r'\int', '∫'),
        (r'\partial', '∂'),
        (r'\,', ' '),
        (r'\;', ' '),
        (r'\!', ''),
        (r'\\', ' ; '),
    ]
    for old, new in pairs:
        s = s.replace(old, new)
    return s


def add_math_runs(paragraph, latex_str):
    """Render simplified LaTeX as runs with subscript/superscript formatting."""
    s = latex_to_unicode(latex_str.strip())
    i = 0
    n = len(s)
    while i < n:
        ch = s[i]
        if ch == '_' and i + 1 < n:
            if s[i + 1] == '{':
                end = find_matching_brace(s, i + 1)
                if end == -1:
                    end = n - 1
                content = s[i + 2:end]
                run = paragraph.add_run(content)
                run.font.subscript = True
                i = end + 1
            else:
                run = paragraph.add_run(s[i + 1])
                run.font.subscript = True
                i += 2
        elif ch == '^' and i + 1 < n:
            if s[i + 1] == '{':
                end = find_matching_brace(s, i + 1)
                if end == -1:
                    end = n - 1
                content = s[i + 2:end]
                run = paragraph.add_run(content)
                run.font.superscript = True
                i = end + 1
            else:
                run = paragraph.add_run(s[i + 1])
                run.font.superscript = True
                i += 2
        elif ch in '{}':
            i += 1
        else:
            j = i
            while j < n and s[j] not in '_^{}':
                j += 1
            paragraph.add_run(s[i:j])
            i = j


def add_runs_with_inline(paragraph, text):
    """Handle **bold**, *italic*, `code`, and inline LaTeX math \\(...\\)."""
    pattern = re.compile(r'(\*\*.+?\*\*|`.+?`|\\\(.+?\\\))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
        elif part.startswith('\\(') and part.endswith('\\)'):
            add_math_runs(paragraph, part[2:-2])
        else:
            paragraph.add_run(part)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        # Skip separator row like |---|---|
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def set_two_columns(section, gap_twips=720):
    """Set a section to 2-column layout. gap_twips: gap between columns in twips (680 ≈ 0.47")."""
    sectPr = section._sectPr
    # Remove existing cols
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), str(gap_twips))
    sectPr.append(cols)


def set_one_column(section):
    """Explicitly set section to 1 column."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)


def convert():
    with open(SRC, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    doc = Document()

    # Page setup: A4 with IEEE-like margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    set_one_column(section)  # Title + abstract span full width

    # Default font: Times New Roman 13pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Times New Roman')
    # Paragraph spacing
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(15)

    i = 0
    in_code_block = False
    code_buffer = []
    in_two_col = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Column-break marker: toggles between 1-col and 2-col
        if stripped == '<!-- columns -->':
            new_section = doc.add_section()
            new_section.page_height = Cm(29.7)
            new_section.page_width = Cm(21.0)
            new_section.left_margin = Cm(1.9)
            new_section.right_margin = Cm(1.9)
            new_section.top_margin = Cm(1.8)
            new_section.bottom_margin = Cm(1.8)
            if not in_two_col:
                set_two_columns(new_section)
                in_two_col = True
            else:
                set_one_column(new_section)
                in_two_col = False
            i += 1
            continue

        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                # close: emit collected code
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Display math block: \[ ... \]
        if stripped == r'\[':
            math_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != r'\]':
                math_lines.append(lines[i].strip())
                i += 1
            i += 1  # skip closing \]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_math_runs(p, ' '.join(math_lines))
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '808080')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(22)
            i += 1
            continue
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].upper())
            run.bold = True
            run.font.size = Pt(14)
            run.font.small_caps = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            i += 1
            continue
        if stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.italic = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Bold abstract / keywords prefix line starting with **
        if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
            run.italic = True
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
            rows, new_i = parse_table(lines, i)
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = ''
                        para = cell.paragraphs[0]
                        add_runs_with_inline(para, cell_text)
                        if r_idx == 0:
                            for r in para.runs:
                                r.bold = True
                        set_cell_border(cell)
                doc.add_paragraph()  # spacing after table
                i = new_i
                continue

        # Bullet list
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list (e.g., "1. ")
        m_num = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            add_runs_with_inline(p, m_num.group(2))
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        # Strip surrounding whitespace, preserve inline formatting
        add_runs_with_inline(p, line)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 2:
        DST = sys.argv[2]
    elif len(sys.argv) > 1:
        SRC = sys.argv[1]
    convert()
